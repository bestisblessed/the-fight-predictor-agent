import openai
import time
import requests
import os
from dotenv import load_dotenv
from PIL import Image
import io
from docx import Document
import subprocess
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
import pickle
import os.path
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")
assistant_mma_handicapper = 'asst_zahT75OFBs5jgi346C9vuzKa' 
if not openai.api_key:
    print("API key is required to run the chatbot.")
    exit()
print("MMA AI Chatbot Initialized - Processing Tweets.")
client = openai.OpenAI(api_key=openai.api_key)
os.makedirs('data', exist_ok=True)
os.makedirs('responses', exist_ok=True)
os.makedirs('files', exist_ok=True)
thread_id = None
tweets_file = 'data/TheFightAgentMentions.docx'
document = Document(tweets_file)
tweets = []
tweet_data = {}  
processed_ids = set()
try:
    with open('data/processed_tweet_ids.txt', 'r') as f:
        processed_ids = {line.strip() for line in f if line.strip()}
except FileNotFoundError:
    print("No previous tweet ID log found, starting fresh")
tweets_file = 'data/TheFightAgentMentions.docx'
document = Document(tweets_file)
tweets = []
tweet_data = {}  
current_tweet = None
current_id = None
temp_tweet = None  
for paragraph in document.paragraphs:
    text = paragraph.text.strip()
    if text.startswith('-----------------------------------'):
        current_id = None
        temp_tweet = None  
    elif text.startswith('Tweet:'):
        temp_tweet = text.replace('Tweet:', '').strip()  
    elif text.startswith('Link:'):
        current_id = text.split('/')[-1].strip()
        print(f"Found Link ID: {current_id}")
        if temp_tweet and current_id and current_id not in processed_ids:
            tweets.append(temp_tweet)
            tweet_data[temp_tweet] = current_id
            print(f"Added tweet with ID: {current_id}")
print(f"\nFound new tweets: {len(tweets)}")
if not tweets:
    print("No new tweets found to process.")
    exit()
print(f"Found {len(tweets)} new tweets to process.")

def get_google_sheets_service():
    SCOPES = [
        'https://www.googleapis.com/auth/spreadsheets',
        # 'https://www.googleapis.com/auth/spreadsheets.readonly',
        'https://www.googleapis.com/auth/drive',
        'https://www.googleapis.com/auth/drive.file'
    ]
    creds = None
    
    if os.path.exists('credentials/token.json'):
        creds = Credentials.from_authorized_user_file('credentials/token.json', SCOPES)
            
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                'credentials/credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        # Save the credentials for the next run
        with open('credentials/token.json', 'w') as token:
            token.write(creds.to_json())
    
    return build('sheets', 'v4', credentials=creds)

SPREADSHEET_ID = '1ojtQSsgGk2hzBeSmxsv_Q1phgWzN0SQbwiKPFkkOYic'
service = get_google_sheets_service()

for tweet in tweets:
    print(f"\nTweet: {tweet}")
    tweet_id = tweet_data[tweet]
    
    # Get AI response (either from existing or generate new)
    try:
        # Check if response exists in spreadsheet
        result = service.spreadsheets().values().get(
            spreadsheetId=SPREADSHEET_ID,
            range=f'Sheet1!A:B'
        ).execute()
        
        values = result.get('values', [])
        existing_response = None
        for row in values:
            if len(row) > 0 and row[0] == tweet_id:
                existing_response = row[1]
                break
                
        if existing_response:
            print(f"Response for tweet ID {tweet_id} already exists. Using existing response.")
            ai_response = existing_response
        else:
            # Generate new response using OpenAI
            thread = client.beta.threads.create()
            thread_id = thread.id
            print(f"New conversation started with Thread ID: {thread_id}")
            message = client.beta.threads.messages.create(
                thread_id=thread_id,
                role="user",
                content=tweet
            )
            run = client.beta.threads.runs.create(
                thread_id=thread_id,
                assistant_id=assistant_mma_handicapper
            )
            print("Processing...")
            while run.status != "completed":
                time.sleep(2)
                run = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run.id)
            messages = client.beta.threads.messages.list(thread_id=thread_id)
            ai_response = None
            for msg in messages.data:
                if msg.role == "assistant":
                    for content in msg.content:
                        if hasattr(content, 'text'):
                            ai_response = content.text.value.strip()
                            print(f"Extracted AI Response: {ai_response}")
                            break
                    if ai_response:
                        break
                        
            if not ai_response:
                print("No valid AI response found in messages.")
                continue  # Skip to next tweet if no valid response
                
            # Only proceed with spreadsheet writing if we have a valid response
            if ai_response:
                # Save response to local file
                response_file_path = f'responses/{tweet_id}.txt'
                try:
                    with open(response_file_path, 'w', encoding='utf-8') as f:
                        f.write(ai_response)
                    print(f"Saved response to {response_file_path}")
                    
                    # Write to spreadsheet
                    values = [[tweet_id, ai_response]]
                    body = {'values': values}
                    service.spreadsheets().values().append(
                        spreadsheetId=SPREADSHEET_ID,
                        range='Sheet1!A:B',
                        valueInputOption='RAW',
                        body=body
                    ).execute()
                except Exception as e:
                    print(f"Error saving response file: {e}")
                    continue
            
        # Use the response to reply on Twitter
        success = False
        try:
            result = subprocess.run(
                ["python", "append_response_to_spreadsheet_in_drive.py", tweet_id],
                check=True,
                capture_output=True,
                text=True
            )
            print(f"Reply script output: {result.stdout}")
            if "Too Many Requests" in result.stdout:
                print("Failed to send reply - Rate limit exceeded")
            else:
                print("Successfully sent reply to Twitter")
                success = True
        except subprocess.CalledProcessError as e:
            print(f"Failed to send reply to Twitter. Error code: {e.returncode}")
            print(f"Error output: {e.stderr}")
            print(f"Standard output: {e.stdout}")

        if success:
            with open('data/processed_tweet_ids.txt', 'a') as f:
                f.write(f"{tweet_id}")
                f.write("")  
                f.write("\n")  
            print(f"Logged processed tweet ID: {tweet_id}")
    except Exception as e:
        print(f"Error processing tweet: {e}")
        continue
    
    time.sleep(5)
