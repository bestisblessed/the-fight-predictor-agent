import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import shutil
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

shutil.copy('../mma-ai/Scrapers/data/fighter_info.csv', 'data/')
shutil.copy('../mma-ai/Scrapers/data/event_data_sherdog.csv', 'data/')

file_path = 'data/fighter_info.csv'
fighter_data = pd.read_csv(file_path)

fighter1_name = 'Dan Hooker'
fighter2_name = 'Justin Gaethje'

def get_fighter_info(name):
    name = name.lower()  
    fighter_info = fighter_data[fighter_data['fighter'] == name]
    if fighter_info.empty:
        return None
    return fighter_info.iloc[0]
fighter1_info = get_fighter_info(fighter1_name)
fighter2_info = get_fighter_info(fighter2_name)

if fighter1_info is None or fighter2_info is None:
    raise ValueError("One or both fighters not found in the dataset.")
fighter1_stats = {
    'Weight Class': fighter1_info['weight class'],
    'Wins': fighter1_info['wins'],
    'Losses': fighter1_info['losses'],
    'Win Streak': fighter1_info['current_win_streak'],
    'Recent Win Rate (5 fights)': fighter1_info['recent_win_rate_5fights'],
    'Height': fighter1_info['height'],
    'Reach': fighter1_info['reach']
}
fighter2_stats = {
    'Weight Class': fighter2_info['weight class'],
    'Wins': fighter2_info['wins'],
    'Losses': fighter2_info['losses'],
    'Win Streak': fighter2_info['current_win_streak'],
    'Recent Win Rate (5 fights)': fighter2_info['recent_win_rate_5fights'],
    'Height': fighter2_info['height'],
    'Reach': fighter2_info['reach']
}

# Create the document
doc = Document()
doc.add_heading(f"Tale of the Tape: {fighter1_name} vs. {fighter2_name}", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Fight Overview
doc.add_heading("Fight Overview", level=2).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph(f"Weight Class: {fighter1_stats['Weight Class']} Bout").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph()

# Fighter Comparison
doc.add_heading("Fighter Comparison\n", level=2).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
comparison_table = doc.add_table(rows=1, cols=3)
comparison_table.style = 'Table Grid'  # Add borders
comparison_table.autofit = False
comparison_table.allow_autofit = False
comparison_table.columns[0].width = Pt(20)
comparison_table.columns[1].width = Pt(120)
comparison_table.columns[2].width = Pt(120)
header_cells = comparison_table.rows[0].cells
header_cells[0].text = ""
header_cells[1].text = fighter1_name
header_cells[2].text = fighter2_name
for cell in header_cells:
    cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))))  # Light gray background
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.bold = True
    run.font.size = Pt(11)
for metric in fighter1_stats.keys():
    row_cells = comparison_table.add_row().cells
    row_cells[0].text = metric
    row_cells[1].text = str(fighter1_stats[metric])
    row_cells[2].text = str(fighter2_stats[metric])
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    for cell in row_cells:
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

# Save the document
output_path = f'reports/tott_{fighter1_name}_{fighter2_name}.docx'
doc.save(output_path)
print(f"Tale of the Tape report saved to {output_path}")


# Working with ufc_database.db
# import sqlite3
# import pandas as pd
# from docx import Document
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.shared import Pt
# import shutil
# shutil.copy('../mma-ai/Scrapers/ufc_database.db', 'data/')
# shutil.copy('../mma-ai/Scrapers/data/fighter_info.csv', 'data/')
# shutil.copy('../mma-ai/Scrapers/data/event_data_sherdog.csv', 'data/')
# db_path = 'data/ufc_database.db'
# conn = sqlite3.connect(db_path)
# fighter1_name = 'Dan Hooker'
# fighter2_name = 'Justin Gaethje'
# recent_fight_query = f"""
#     SELECT *
#     FROM fight_results
#     WHERE FIGHTER1 IN ('{fighter1_name}', '{fighter2_name}')
#        OR FIGHTER2 IN ('{fighter1_name}', '{fighter2_name}')
#     ORDER BY DATE DESC
#     LIMIT 2;
# """
# recent_fights = pd.read_sql(recent_fight_query, conn)
# fighter1_recent = recent_fights[
#     (recent_fights['FIGHTER1'] == fighter1_name) | (recent_fights['FIGHTER2'] == fighter1_name)
# ]
# fighter2_recent = recent_fights[
#     (recent_fights['FIGHTER1'] == fighter2_name) | (recent_fights['FIGHTER2'] == fighter2_name)
# ]
# fighter1_weightclass = (
#     fighter1_recent['WEIGHTCLASS'].iloc[0] if not fighter1_recent.empty else "Unknown"
# )
# fighter2_weightclass = (
#     fighter2_recent['WEIGHTCLASS'].iloc[0] if not fighter2_recent.empty else "Unknown"
# )
# win_loss_query = f"""
#     SELECT WINNING_FIGHTER AS FIGHTER, COUNT(*) AS WINS
#     FROM fight_results
#     WHERE WINNING_FIGHTER IN ('{fighter1_name}', '{fighter2_name}')
#     GROUP BY WINNING_FIGHTER;
# """
# wins_data = pd.read_sql(win_loss_query, conn)
# total_fights_query = f"""
#     SELECT FIGHTER1 AS FIGHTER, COUNT(*) AS TOTAL_FIGHTS
#     FROM fight_results
#     WHERE FIGHTER1 IN ('{fighter1_name}', '{fighter2_name}')
#     GROUP BY FIGHTER1
#     UNION ALL
#     SELECT FIGHTER2 AS FIGHTER, COUNT(*) AS TOTAL_FIGHTS
#     FROM fight_results
#     WHERE FIGHTER2 IN ('{fighter1_name}', '{fighter2_name}')
#     GROUP BY FIGHTER2;
# """
# total_fights_data = pd.read_sql(total_fights_query, conn)
# total_fights_agg = total_fights_data.groupby('FIGHTER', as_index=False).sum()
# summary = total_fights_agg.merge(wins_data, on='FIGHTER', how='left')
# summary['LOSSES'] = summary['TOTAL_FIGHTS'] - summary['WINS'].fillna(0)

# # Create the document
# doc = Document()
# doc.add_heading(f"Tale of the Tape: {fighter1_name} vs. {fighter2_name}", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# # Fight Overview
# overview = doc.add_heading("Fight Overview", level=2)
# overview.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# doc.add_paragraph().alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
# doc.add_paragraph(f"Event: Upcoming UFC Event").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# doc.add_paragraph(f"Weight Class: {fighter1_weightclass} Bout").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# doc.add_paragraph().alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 

# # Fighter Comparison
# fighter_comparison = doc.add_heading("Fighter Comparison", level=2)
# fighter_comparison.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# comparison_table = doc.add_table(rows=1, cols=3)
# comparison_table.autofit = False  
# comparison_table.columns[0].width = Pt(100)  
# comparison_table.columns[1].width = Pt(200)  
# comparison_table.columns[2].width = Pt(200)  
# doc.add_paragraph() 
# name_cells = comparison_table.add_row().cells
# name_cells[1].text = fighter1_name
# name_cells[2].text = fighter2_name
# for metric, f1_data, f2_data in [
#     ("Total Fights", summary.loc[summary['FIGHTER'] == fighter1_name, 'TOTAL_FIGHTS'].values[0],
#      summary.loc[summary['FIGHTER'] == fighter2_name, 'TOTAL_FIGHTS'].values[0]),
#     ("Wins", summary.loc[summary['FIGHTER'] == fighter1_name, 'WINS'].values[0],
#      summary.loc[summary['FIGHTER'] == fighter2_name, 'WINS'].values[0]),
#     ("Losses", summary.loc[summary['FIGHTER'] == fighter1_name, 'LOSSES'].values[0],
#      summary.loc[summary['FIGHTER'] == fighter2_name, 'LOSSES'].values[0]),
# ]:
#     row_cells = comparison_table.add_row().cells
#     row_cells[0].text = metric
#     row_cells[1].text = str(f1_data)
#     row_cells[2].text = str(f2_data)
# for row in comparison_table.rows:
#     for cell in row.cells:
#         for paragraph in cell.paragraphs:
#             paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#             paragraph.style.font.size = Pt(10)

# # Save the document
# output_path = f'reports/tott_{fighter1_name}_{fighter2_name}.docx'
# doc.save(output_path)
# print(f"Tale of the Tape report saved to {output_path}")
