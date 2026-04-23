import openai
import time
import requests
import os
from dotenv import load_dotenv
from PIL import Image
import io
from docx import Document
import subprocess
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")
assistant_mma_handicapper = 'asst_zahT75OFBs5jgi346C9vuzKa' 
if not openai.api_key:
    print("API key is required to run the chatbot.")
    exit()
print("MMA AI Chatbot initialized. Processing tweets from document.")
client = openai.OpenAI(api_key=openai.api_key)
os.makedirs('data', exist_ok=True)
thread_id = None
tweets_file = 'data/TheFightAgentMentions.docx'
document = Document(tweets_file)
tweets = []
tweet_data = {}  # Store both tweet text and ID together

# After initial imports and before processing tweets, load processed IDs
processed_ids = set()
try:
    with open('data/processed_tweet_ids.txt', 'r') as f:
        processed_ids = {line.strip() for line in f if line.strip()}
    print(f"Loaded {len(processed_ids)} processed tweet IDs")
except FileNotFoundError:
    print("No previous tweet ID log found, starting fresh")

print(f"\nDebug: Loaded processed IDs: {processed_ids}")

# Path to the docx file containing tweets
tweets_file = 'data/TheFightAgentMentions.docx'
document = Document(tweets_file)
tweets = []
tweet_data = {}  # Store both tweet text and ID together

# Gather tweets that haven't been processed yet
current_tweet = None
current_id = None
temp_tweet = None  # Store tweet text temporarily until we find its ID

for paragraph in document.paragraphs:
    text = paragraph.text.strip()
    if text.startswith('-----------------------------------'):
        current_id = None
        temp_tweet = None  # Reset temporary tweet storage
    elif text.startswith('Tweet:'):
        temp_tweet = text.replace('Tweet:', '').strip()  # Store tweet text
    elif text.startswith('Link:'):
        current_id = text.split('/')[-1].strip()
        print(f"Found Link ID: {current_id}")
        # Now that we have both tweet and ID, check if we should process it
        if temp_tweet and current_id and current_id not in processed_ids:
            tweets.append(temp_tweet)
            tweet_data[temp_tweet] = current_id
            print(f"Added tweet with ID: {current_id}")

print(f"\nDebug: Found tweets: {len(tweets)}")

if not tweets:
    print("No new tweets found to process.")
    exit()

print(f"Found {len(tweets)} new tweets to process.")

for tweet in tweets:
    print(f"\nTweet: {tweet}")
    
    # Create a new thread for each tweet
    thread = client.beta.threads.create()
    thread_id = thread.id
    print(f"New conversation started with Thread ID: {thread_id}")

    message = client.beta.threads.messages.create(
        thread_id=thread_id,
        role="user",
        content=tweet
    )

    run = client.beta.threads.runs.create(
        thread_id=thread_id,
        assistant_id=assistant_mma_handicapper
    )
    print("Processing...")
    while run.status != "completed":
        time.sleep(2)
        run = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run.id)
    messages = client.beta.threads.messages.list(thread_id=thread_id)
    # Skip the first message (user's question) by checking message role
    for message in reversed(messages.data):
        if hasattr(message.content[0], 'text'):
            ai_response = message.content[0].text.value
            # Skip if this is the user's original question
            if message.role == "user":
                continue
            print(f"AI: {ai_response}")
            tweet_id = tweet_data[tweet]
            
            # # Save the tweet and response to a file named after the tweet ID
            # timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
            # response_file = f'responses/{tweet_id}.txt'
            # with open(response_file, 'w', encoding='utf-8') as f:
            #     f.write(f"Timestamp: {timestamp}\n")
            #     f.write(f"Tweet ID: {tweet_id}\n")
            #     f.write(f"Original Tweet: {tweet}\n")
            #     f.write(f"AI Response: {ai_response}\n")
            
            # print(f"Saved response to {response_file}")
            
            # Call the reply script as a subprocess with enhanced error handling
            try:
                # Get the full Python path
                python_path = subprocess.run(['which', 'python'], 
                    capture_output=True, 
                    text=True, 
                    check=True
                ).stdout.strip()
                
                result = subprocess.run(
                    [python_path, 'X/reply_single_tweet.py', tweet_id, ai_response], 
                    check=True,
                    capture_output=True,
                    text=True
                )
                print(f"Reply script output: {result.stdout}")
                
                # If rate limited, skip EVERYTHING else
                if "429 Too Many Requests" in result.stdout:
                    print("Rate limit hit, will try this tweet again later")
                    exit()  # Exit the entire program if rate limited
                
                # Everything below this only happens if NO rate limit hit
                print("Successfully sent reply to Twitter")
                
                os.makedirs('responses', exist_ok=True)
                response_file = f'responses/{tweet_id}.txt'
                with open(response_file, 'w', encoding='utf-8') as f:
                    f.write(ai_response)
                print(f"Saved response to {response_file}")
                
                # Upload to Drive
                try:
                    python_path = subprocess.run(['which', 'python'], 
                        capture_output=True, 
                        text=True, 
                        check=True
                    ).stdout.strip()
                    
                    result = subprocess.run(
                        [python_path, 'X/upload_responses_to_drive.py', tweet_id],
                        check=True,
                        capture_output=True,
                        text=True
                    )
                    print(f"Drive upload output: {result.stdout}")
                    
                    # Mark as processed
                    with open('data/processed_tweet_ids.txt', 'a') as f:
                        f.write(f"{tweet_id}\n")
                    print(f"Logged processed tweet ID: {tweet_id}")
                    
                    # Remove from document
                    paragraphs_to_keep = []
                    in_block_to_remove = False
                    separator_count = 0
                    
                    for paragraph in document.paragraphs:
                        text = paragraph.text.strip()
                        
                        # Count separators to track block boundaries
                        if text.startswith('-----------------------------------'):
                            separator_count += 1
                            # If we're at the end of a block we want to remove, skip this separator
                            if in_block_to_remove and separator_count % 2 == 0:
                                in_block_to_remove = False
                                continue
                            # If we're not removing this block, keep the separator
                            if not in_block_to_remove:
                                paragraphs_to_keep.append(text)
                            continue

                        # Check if this is the tweet we want to remove
                        if text.startswith('Tweet:') and text.replace('Tweet:', '').strip() == tweet:
                            in_block_to_remove = True
                            continue

                        # Only keep text if we're not in a block to remove
                        if not in_block_to_remove:
                            paragraphs_to_keep.append(text)
                    
                    # Create new document with remaining content
                    new_doc = Document()
                    for text in paragraphs_to_keep:
                        if text:  # Only add non-empty paragraphs
                            new_doc.add_paragraph(text)
                    
                    new_doc.save(tweets_file)
                    document = new_doc
                    print("Removed processed tweet and related content from document")
                    
                except subprocess.CalledProcessError as e:
                    print(f"Failed to upload to Drive. Error: {e.stderr}")
                    exit()
                    
            except subprocess.CalledProcessError as e:
                print(f"Failed to send reply to Twitter. Error code: {e.returncode}")
                print(f"Error output: {e.stderr}")
                print(f"Standard output: {e.stdout}")
                exit()
        elif hasattr(message.content[0], 'image_file'):
            print("AI: [Image file received]")
            file_id = message.content[0].image_file.file_id
            file_url = f"https://api.openai.com/v1/files/{file_id}/content"
            headers = {"Authorization": f"Bearer {openai.api_key}"}
            print("Downloading image...")
            image_data = requests.get(file_url, headers=headers)
            if image_data.status_code == 200:
                filename = f"data/assistant_image_{int(time.time())}.png"
                with open(filename, "wb") as f:
                    f.write(image_data.content)
                print(f"Image saved {filename}")
                img = Image.open(filename)
                img.show()  
            else:
                print("Failed to download the image.")
        else:
            print("AI: [Unsupported content type]")

    # After successful processing, log the tweet ID and remove the block
    tweet_id = tweet_data[tweet]  # Get the ID for this tweet
    
    # Save the raw AI response to a text file
    os.makedirs('responses', exist_ok=True)
    response_file = f'responses/{tweet_id}.txt'
    with open(response_file, 'w', encoding='utf-8') as f:
        f.write(ai_response)  # Just save the raw AI response
    print(f"Saved response to {response_file}")
    
    # Upload the response file to Google Drive
    try:
        # Get the full Python path
        python_path = subprocess.run(['which', 'python'], 
            capture_output=True, 
            text=True, 
            check=True
        ).stdout.strip()
        
        result = subprocess.run(
            [python_path, 'X/upload_responses_to_drive.py', tweet_id],
            check=True,
            capture_output=True,
            text=True
        )
        print(f"Drive upload output: {result.stdout}")
        
        # Only mark as processed if both upload and tweet reply were successful
        if "429 Too Many Requests" not in result.stdout:
            with open('data/processed_tweet_ids.txt', 'a') as f:
                f.write(f"{tweet_id}\n")
            print(f"Logged processed tweet ID: {tweet_id}")
        else:
            print(f"Rate limit hit for tweet {tweet_id}, will try again later")
            continue  # Skip to next tweet without marking this one as processed
            
    except subprocess.CalledProcessError as e:
        print(f"Failed to upload to Drive. Error: {e.stderr}")
        continue  # Skip to next tweet without marking this one as processed

    # Remove the processed tweet from document
    paragraphs_to_keep = []
    in_block_to_remove = False
    separator_count = 0
    
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        
        # Count separators to track block boundaries
        if text.startswith('-----------------------------------'):
            separator_count += 1
            # If we're at the end of a block we want to remove, skip this separator
            if in_block_to_remove and separator_count % 2 == 0:
                in_block_to_remove = False
                continue
            # If we're not removing this block, keep the separator
            if not in_block_to_remove:
                paragraphs_to_keep.append(text)
            continue

        # Check if this is the tweet we want to remove
        if text.startswith('Tweet:') and text.replace('Tweet:', '').strip() == tweet:
            in_block_to_remove = True
            continue

        # Only keep text if we're not in a block to remove
        if not in_block_to_remove:
            paragraphs_to_keep.append(text)
    
    # Create new document with remaining content
    new_doc = Document()
    for text in paragraphs_to_keep:
        if text:  # Only add non-empty paragraphs
            new_doc.add_paragraph(text)
    
    new_doc.save(tweets_file)
    document = new_doc
    print("Removed processed tweet and related content from document")


